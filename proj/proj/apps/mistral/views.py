from django.shortcuts import render
from django.http import JsonResponse, HttpResponse
from mistralai import Mistral
import os
from docx import Document
from docx.shared import Pt
import tempfile
import markdown
import re
from django.contrib.auth.decorators import login_required
from home.models import UserStatistics

# Home view to render the interface
@login_required
def main(request):
    if request.method == "POST":
        prompt = request.POST.get("prompt")
        response = get_mistral_response(prompt)
        request.session['response_text'] = response  # Save raw response in session for download
        if request.user.is_authenticated:
            user_statistics = UserStatistics.objects.get(user=request.user)
            user_statistics.successful_text_analyses += 1
            user_statistics.save()
        return JsonResponse({"response": response})
    return render(request, "mistral_index.html")

# Function to get response from Mistral API
def get_mistral_response(prompt):
    system_prompt = "Ты — помощник, который определяет, какие характеристики нужно указывать у выбранного устройства для размещения заказа. На выходе выдаешь опросный лист, где достаточно вставить нужные параметры или отметить галочку в чекбоксе. Без приветствий в начале и заключений в конце"
    api_key = os.getenv("MISTRAL_API_KEY")
    if not api_key:
        return "API-ключ Mistral не установлен."

    client = Mistral(api_key=api_key)
    try:
        response = client.chat.complete(
            model="mistral-large-latest",
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": prompt}
            ]
        )
        return response.choices[0].message.content
    except Exception as e:
        return f"Ошибка при обращении к Mistral API: {str(e)}"

# Function to convert Markdown to Word-compatible format
def markdown_to_word(markdown_text, output_file):
    # Create a new document
    doc = Document()

    # Split text into lines
    lines = markdown_text.split('\n')

    for line in lines:
        line = line.strip()
        if not line:
            continue

        # Process headings
        if line.startswith('# '):
            doc.add_heading(line[2:].strip(), level=1)
        elif line.startswith('## '):
            doc.add_heading(line[3:].strip(), level=2)
        elif line.startswith('### '):
            doc.add_heading(line[4:].strip(), level=3)
        # Process bold text
        elif '**' in line:
            paragraph = doc.add_paragraph()
            parts = re.split(r'(\*\*.*?\*\*)', line)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    run = paragraph.add_run(part[2:-2].strip())
                    run.bold = True
                else:
                    paragraph.add_run(part)
        # Process italic text
        elif '*' in line:
            paragraph = doc.add_paragraph()
            parts = re.split(r'(\*.*?\*)', line)
            for part in parts:
                if part.startswith('*') and part.endswith('*'):
                    run = paragraph.add_run(part[1:-1].strip())
                    run.italic = True
                else:
                    paragraph.add_run(part)
        # Process underline text
        elif '__' in line:
            paragraph = doc.add_paragraph()
            parts = re.split(r'(\_\_.*?\_\_)', line)
            for part in parts:
                if part.startswith('__') and part.endswith('__'):
                    run = paragraph.add_run(part[2:-2].strip())
                    run.underline = True
                else:
                    paragraph.add_run(part)
        # Process blockquote
        elif line.startswith('> '):
            paragraph = doc.add_paragraph(line[2:].strip())
            paragraph_format = paragraph.paragraph_format
            paragraph_format.left_indent = Pt(36)  # Indent from the left
        # Process unordered list
        elif line.startswith('- '):
            doc.add_paragraph(line[2:].strip(), style='List Bullet')
        # Process ordered list
        elif line[0].isdigit() and line[1:3] == '. ' and len(line) > 3:
            doc.add_paragraph(line[3:].strip(), style='List Number')
        # Regular text
        else:
            doc.add_paragraph(line)

    # Save document
    doc.save(output_file)

# Function to download response as a Word document
@login_required
def download_docx(request):
    response_text = request.session.get('response_text', None)
    if not response_text:
        return HttpResponse("Нет текста для скачивания.", status=400)

    with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
        markdown_to_word(response_text, tmp.name)
        tmp.seek(0)
        with open(tmp.name, 'rb') as f:
            response = HttpResponse(f.read(), content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
            response['Content-Disposition'] = 'attachment; filename=Mistral_Response.docx'
            return response



    # system_prompt = "Ты — помощник, который определяет, какие характеристики нужно указывать у выбранного устройства для размещения заказа. На выходе выдаешь опросный лист, где достаточно вставить нужные параметры или отметить галочку в чекбоксе. Без приветствий в начале и заключений в конце"
